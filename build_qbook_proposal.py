from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from pathlib import Path


OUT = Path('deliverables') / 'Qbook_Internal_Facility_Booking_and_Approval_System_Proposal.docx'

NAVY = '0B2545'
BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
MUTED = '5B6573'
LIGHT_BLUE = 'E8EEF5'
LIGHT_GRAY = 'F4F6F9'
BORDER = 'C9D3DE'
WHITE = 'FFFFFF'
BLACK = '000000'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = 'w:{}'.format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ['val', 'sz', 'space', 'color']:
                if key in kwargs[edge]:
                    element.set(qn('w:{}'.format(key)), str(kwargs[edge][key]))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement('w:keepNext')
    p_pr.append(keep)


def set_run_font(run, size=None, color=BLACK, bold=None, italic=None):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in [
        ('Heading 1', 16, BLUE, 18, 10),
        ('Heading 2', 13, BLUE, 12, 6),
        ('Heading 3', 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    toc = doc.styles.add_style('TOC Entry', WD_STYLE_TYPE.PARAGRAPH)
    toc.font.name = 'Calibri'
    toc._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    toc._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    toc.font.size = Pt(11)
    toc.font.color.rgb = RGBColor.from_string(NAVY)
    toc.paragraph_format.space_after = Pt(5)
    toc.paragraph_format.line_spacing = 1.15

    label = doc.styles.add_style('Metadata Label', WD_STYLE_TYPE.PARAGRAPH)
    label.font.name = 'Calibri'
    label.font.size = Pt(10)
    label.font.bold = True
    label.font.color.rgb = RGBColor.from_string(MUTED)
    label.paragraph_format.space_after = Pt(1)

    small = doc.styles.add_style('Small Note', WD_STYLE_TYPE.PARAGRAPH)
    small.font.name = 'Calibri'
    small.font.size = Pt(9)
    small.font.color.rgb = RGBColor.from_string(MUTED)
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.0


def setup_section(section, first=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    if not first:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.paragraph_format.space_after = Pt(0)
        run = header.add_run('QBOOK  |  INTERNAL FACILITY BOOKING AND APPROVAL SYSTEM PROPOSAL')
        set_run_font(run, size=8.5, color=MUTED, bold=True)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.paragraph_format.space_before = Pt(0)
        footer.paragraph_format.space_after = Pt(0)
        run = footer.add_run('Qhazanah Sabah Berhad  |  Internal Use  |  Page ')
        set_run_font(run, size=8.5, color=MUTED)
        add_page_number(footer)


def add_paragraph(doc, text='', bold_prefix=None, after=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_section_heading(doc, number, title):
    p = doc.add_paragraph(style='Heading 1')
    r = p.add_run(f'{number}. {title}')
    set_run_font(r, size=16, color=BLUE, bold=True)
    set_keep_with_next(p)
    return p


def add_toc_entry(doc, number, title, page):
    p = doc.add_paragraph(style='TOC Entry')
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(f'{number}. {title}')
    set_run_font(r, size=11, color=NAVY)
    r = p.add_run('\t' + str(page))
    set_run_font(r, size=11, color=NAVY)
    return p


def add_role_table(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [1900, 7460])
    table.style = 'Table Grid'
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, ['Role', 'Core responsibilities']):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=NAVY, bold=True)
    data = [
        ('Employee', [
            'Manage own bookings, including changes, rescheduling, and cancellations.',
            'Invite internal attendees and tag involved departments.',
            'View booking invitations, RSVP status, booking status, and calendar information.',
        ]),
        ('Admin', [
            'Review, approve, or reject bookings in accordance with defined requirements.',
            'Manage operational booking records and provide day-to-day support.',
            'Review booking activity, reports, and notification status.',
        ]),
        ('Super Admin', [
            'Manage users, roles, facilities, departments, department mailboxes, settings, reports, notifications, and audit records.',
        ]),
    ]
    for role, items in data:
        row = table.add_row()
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        row.cells[0].text = ''
        p = row.cells[0].paragraphs[0]
        r = p.add_run(role)
        set_run_font(r, size=10.5, color=NAVY, bold=True)
        row.cells[1].text = ''
        for idx, item in enumerate(items):
            p = row.cells[1].paragraphs[0] if idx == 0 else row.cells[1].add_paragraph()
            p.style = 'List Bullet'
            p.paragraph_format.left_indent = Inches(0.24)
            p.paragraph_format.first_line_indent = Inches(-0.16)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(item)
            set_run_font(r, size=9.7)
        for cell in row.cells:
            set_cell_border(cell, top={'val':'single','sz':'4','color':BORDER}, bottom={'val':'single','sz':'4','color':BORDER}, left={'val':'single','sz':'4','color':BORDER}, right={'val':'single','sz':'4','color':BORDER})
    return table


def add_pilot_table(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2500, 6860])
    table.style = 'Table Grid'
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, ['Pilot area', 'Proposed scope']):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ''
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, size=10.5, color=NAVY, bold=True)
    rows = [
        ('Participants', 'IT/Admin representatives and one selected department, with a defined group of internal staff users.'),
        ('Booking workflow', 'Create bookings; check availability; add department tags; invite internal attendees; submit for approval; approve or reject; issue cancellations.'),
        ('Communication checks', 'Confirm booking confirmations, reminders, approval, rejection, invitation, and cancellation emails reach staff and relevant department mailboxes.'),
        ('Acceptance outcome', 'Record user feedback, confirm operating responsibilities, resolve pilot findings, and obtain acceptance sign-off before recommending wider rollout.'),
    ]
    for left, right in rows:
        row = table.add_row()
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for cell, txt, bold in [(row.cells[0], left, True), (row.cells[1], right, False)]:
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            set_run_font(r, size=10, color=NAVY if bold else BLACK, bold=bold)
            set_cell_border(cell, top={'val':'single','sz':'4','color':BORDER}, bottom={'val':'single','sz':'4','color':BORDER}, left={'val':'single','sz':'4','color':BORDER}, right={'val':'single','sz':'4','color':BORDER})
    return table


def add_approval_table(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2250, 2150, 2200, 2760])
    table.style = 'Table Grid'
    headers = ['Approval item', 'Name', 'Signature', 'Date']
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ''
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, size=10, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    for role in ['Business sponsor / Management representative', 'IT owner / System owner', 'Pilot acceptance representative']:
        row = table.add_row()
        for i, cell in enumerate(row.cells):
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(14)
            r = p.add_run(role if i == 0 else '')
            set_run_font(r, size=9.6, color=NAVY if i == 0 else BLACK, bold=(i == 0))
            set_cell_border(cell, top={'val':'single','sz':'4','color':BORDER}, bottom={'val':'single','sz':'4','color':BORDER}, left={'val':'single','sz':'4','color':BORDER}, right={'val':'single','sz':'4','color':BORDER})
    return table


def build_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    setup_section(doc.sections[0], first=True)
    sec = doc.sections[0]

    # Cover page: proposal_centerpiece pattern.
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('QHAZANAH SABAH BERHAD')
    set_run_font(r, size=13, color=MUTED, bold=True)
    p.paragraph_format.space_after = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Qbook: Internal Facility Booking\nand Approval System Proposal')
    set_run_font(r, size=25, color=NAVY, bold=True)
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Proposal for a controlled internal pilot rollout')
    set_run_font(r, size=13, color=MUTED, italic=True)
    p.paragraph_format.space_after = Pt(26)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(16)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    meta = doc.add_table(rows=3, cols=2)
    set_table_geometry(meta, [4680, 4680])
    for row in meta.rows:
        for cell in row.cells:
            set_cell_shading(cell, LIGHT_GRAY)
            set_cell_border(cell, top={'val':'single','sz':'4','color':WHITE}, bottom={'val':'single','sz':'4','color':WHITE}, left={'val':'single','sz':'4','color':WHITE}, right={'val':'single','sz':'4','color':WHITE})
    meta_data = [
        ('Version', '0.1'), ('Date', '22 July 2026'),
        ('Prepared by', '[Name / Department]'), ('Document status', 'Draft for management and IT review'),
        ('Classification', 'Internal Use'), ('Decision sought', 'Approval to proceed with a controlled pilot'),
    ]
    for cell, (label, value) in zip([c for r in meta.rows for c in r.cells], meta_data):
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(label.upper())
        set_run_font(r, size=8.5, color=MUTED, bold=True)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=NAVY, bold=(label in ['Version', 'Date']))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    r = p.add_run('This proposal outlines the intended Qbook capability, pilot scope, governance controls, and information required before wider rollout consideration.')
    set_run_font(r, size=10.5, color=MUTED, italic=True)

    doc.add_page_break()
    # Contents page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('TABLE OF CONTENTS')
    set_run_font(r, size=18, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(18)
    toc_entries = [
        ('1', 'Executive Summary', 3),
        ('2', 'Background and Current Challenges', 3),
        ('3', 'Proposed Solution', 4),
        ('4', 'User Roles and Responsibilities', 5),
        ('5', 'Department Coordination', 5),
        ('6', 'Administration and Configurability', 6),
        ('7', 'Security and Governance', 6),
        ('8', 'Proposed Pilot Rollout', 7),
        ('9', 'Information Required from the Company and IT Team', 7),
        ('10', 'Decision Requested', 8),
        ('11', 'Conclusion', 8),
        ('', 'Approval and Sign-Off', 8),
    ]
    for number, title, page in toc_entries:
        add_toc_entry(doc, number, title, page)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run('Document purpose')
    set_run_font(r, size=11, color=NAVY, bold=True)
    add_paragraph(doc, 'To seek management and IT alignment on a controlled Qbook pilot and the information needed to set up the proposed service responsibly.', after=0)

    doc.add_page_break()
    # Page 3
    add_section_heading(doc, '1', 'Executive Summary')
    add_paragraph(doc, 'Qbook is a proposed internal web-based facility booking system for Qhazanah Sabah Berhad. It is intended to replace spreadsheet-based room and facility booking processes with a central, traceable platform that brings booking requests, approvals, attendees, departments, and notifications into one controlled workflow.')
    add_paragraph(doc, 'The proposed system would give staff a clearer way to request and manage bookings, while providing authorised administrators with visibility of status, approvals, and key booking activity. By consolidating information that is currently handled across Excel files and manual coordination, Qbook is intended to improve reliability, reduce avoidable conflicts, and support more consistent operational oversight.')
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), LIGHT_GRAY)
    p_pr.append(shd)
    r = p.add_run('Recommendation: approve a controlled pilot involving IT/Admin and one selected department before any wider rollout decision is made.')
    set_run_font(r, size=11, color=NAVY, bold=True)

    add_section_heading(doc, '2', 'Background and Current Challenges')
    add_paragraph(doc, 'Facility and room bookings are currently managed through Excel-based processes. While spreadsheets can capture booking details, they depend heavily on manual checking, individual coordination, and timely updates by multiple parties. This makes it difficult to maintain a single, dependable view of availability and approval status as booking activity grows.')
    add_bullet(doc, 'Manual availability checking can be slow and may not reflect the latest change at the time a request is made.')
    add_bullet(doc, 'Duplicate, incomplete, or conflicting information can arise when booking details are shared or updated across files and messages.')
    add_bullet(doc, 'Employees and administrators have limited visibility of whether a booking is pending, approved, rejected, changed, or cancelled.')
    add_bullet(doc, 'Approval coordination is manual, which can delay decisions and make follow-up responsibilities unclear.')
    add_bullet(doc, 'Attendees and involved departments are difficult to track consistently across booking records.')
    add_bullet(doc, 'There is no single, clear audit trail or notification history for important booking actions.')
    add_paragraph(doc, 'These limitations create operational risk: staff may act on outdated information, approvers may miss requests, and management may have limited evidence when reviewing booking activity or resolving a dispute.', bold_prefix='These limitations create operational risk:')

    doc.add_page_break()
    # Page 4
    add_section_heading(doc, '3', 'Proposed Solution')
    add_paragraph(doc, 'Qbook is proposed as a central internal system for facility booking and approval management. It would provide a consistent workflow from booking request through approval, communication, and record keeping, while maintaining visibility for the appropriate users and administrators.')
    p = doc.add_paragraph()
    r = p.add_run('The proposed system enables staff to:')
    set_run_font(r, size=11, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(5)
    for item in [
        'Create, edit, reschedule, and cancel bookings within their own permitted access.',
        'Check facility availability and prevent overlapping time conflicts before a booking is confirmed.',
        'Invite internal attendees and track their RSVP responses.',
        'Tag one or more involved departments for a booking.',
        'Notify department mailboxes about relevant booking updates.',
        'Submit bookings for approval where the selected facility requires it.',
        'Receive booking confirmations, reminders, approval, rejection, invitation, and cancellation notifications.',
        'View booking calendars and printable approval forms for operational reference and record keeping.',
    ]:
        add_bullet(doc, item)
    add_paragraph(doc, 'The solution is intended to become the central record for participating facilities during the pilot. It does not remove the need for appropriate operational judgement by administrators or approvers; rather, it provides a clearer system of record for the actions they take.')
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run('Pilot principle: functionality should be validated through real, controlled use before wider deployment is considered.')
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True, italic=True)

    doc.add_page_break()
    # Page 5
    add_section_heading(doc, '4', 'User Roles and Responsibilities')
    add_paragraph(doc, 'Qbook is designed around defined roles so that staff can complete ordinary booking tasks while administrative and governance functions remain controlled. Final role assignments should be confirmed by the company before pilot access is provided.')
    add_role_table(doc)
    add_section_heading(doc, '5', 'Department Coordination')
    add_paragraph(doc, 'A booking may optionally include one or more involved departments, such as HR, IT, CMCD, Facilities, or other departments identified by the company. This allows the booking record to show which operational teams are relevant to an activity without requiring a separate spreadsheet or informal email trail.')
    add_paragraph(doc, 'Each department would be linked to an existing department mailbox supplied by the company. When a department is included in a booking, Qbook is intended to provide clear visibility of that involvement in the booking details, approval flow, printable approval form, and relevant notification emails.')
    add_paragraph(doc, 'Department tags should support coordination rather than replace department-level approval policies. Approval requirements remain governed by the facility configuration and authorised administrative process.', bold_prefix='Department tags should support coordination rather than replace department-level approval policies.')

    doc.add_page_break()
    # Page 6
    add_section_heading(doc, '6', 'Administration and Configurability')
    add_paragraph(doc, 'Authorised administrators would manage the operating configuration required for Qbook to remain accurate and useful. Administrative access should be limited to personnel nominated by the company and reviewed as responsibilities change.')
    for item in [
        'Facilities included in the system and their relevant booking requirements.',
        'Staff access and assigned user roles.',
        'Departments and their mailbox details.',
        'Approval requirements for facilities.',
        'Booking records, reports, email notifications, and audit history.',
    ]:
        add_bullet(doc, item)
    add_paragraph(doc, 'Facilities can be added or updated through the Admin console after rollout, subject to the company’s normal approval and operating procedures. This means the proposed service can begin with an agreed pilot scope and expand over time without requiring the original facility list to remain fixed.')

    add_section_heading(doc, '7', 'Security and Governance')
    add_paragraph(doc, 'The pilot should operate with controls appropriate for an internal company service. The following controls are proposed as the baseline governance model:')
    for item in [
        'Microsoft company-email sign-in for user authentication.',
        'Role-based access control so users see and manage only the functions appropriate to their role.',
        'Internal staff-only access, managed through company-controlled user provisioning.',
        'Booking conflict prevention to reduce competing reservations for the same facility and time.',
        'Approval workflow controls where the selected facility requires approval.',
        'An audit trail for important booking actions, including key status changes and administrative activity.',
        'Controlled access to administrative functions, with Super Admin responsibility for higher-risk configuration and governance tasks.',
    ]:
        add_bullet(doc, item)
    add_paragraph(doc, 'These controls are proposed to support accountability and operational continuity. Detailed implementation settings, including Microsoft login, email, and calendar integration arrangements, should be confirmed with the IT team as part of pilot preparation.')

    doc.add_page_break()
    # Page 7
    add_section_heading(doc, '8', 'Proposed Pilot Rollout')
    add_paragraph(doc, 'It is recommended that Qbook first be introduced through a controlled pilot involving IT/Admin and one selected department. The pilot should use real internal booking scenarios within an agreed, limited facility scope. This will allow the company to validate the process, communications, and role responsibilities before considering wider use.')
    add_pilot_table(doc)

    add_section_heading(doc, '9', 'Information Required from the Company and IT Team')
    add_paragraph(doc, 'To prepare the pilot safely and efficiently, the following information and confirmations are required:')
    for item in [
        'Named system owner, Super Admins, and Admins.',
        'Existing department mailboxes.',
        'Initial staff list: full name, company email, department, and assigned role.',
        'Confirmation of facilities to include at launch.',
        'Approved Qbook sender mailbox/address.',
        'Approval to send booking-related emails to staff and department mailboxes.',
        'Pilot participants and acceptance sign-off.',
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    # Page 8
    add_section_heading(doc, '10', 'Decision Requested')
    add_paragraph(doc, 'Management and IT stakeholders are requested to approve the following next steps:')
    for item in [
        'Proceed with a controlled Qbook pilot rollout.',
        'Nominate the relevant IT/Admin owners.',
        'Provide the required launch information listed in this proposal.',
        'Gather feedback on the proposed Microsoft login, email, and calendar integration structure.',
    ]:
        add_bullet(doc, item)

    add_section_heading(doc, '11', 'Conclusion')
    add_paragraph(doc, 'Qbook provides a proposed structured, transparent, and scalable alternative to Excel-based facility booking management. By centralising bookings, approvals, attendees, department involvement, notifications, and key activity records, it is intended to give staff and administrators a clearer and more accountable way to manage internal facilities.')
    add_paragraph(doc, 'A controlled pilot will allow Qhazanah Sabah Berhad to validate the practical workflow and governance model before a wider rollout decision. Subject to successful pilot acceptance, the company can expand the system over time by adding facilities, departments, and authorised users through the Admin console.')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15)
    r = p.add_run('Approval and Sign-Off')
    set_run_font(r, size=13, color=BLUE, bold=True)
    p.paragraph_format.space_after = Pt(6)
    add_paragraph(doc, 'The signatories below acknowledge review of this proposal and the decision on whether to proceed with the controlled pilot.', after=7)
    add_approval_table(doc)

    doc.core_properties.title = 'Qbook: Internal Facility Booking and Approval System Proposal'
    doc.core_properties.subject = 'Internal proposal for controlled Qbook pilot rollout'
    doc.core_properties.author = '[Name / Department]'
    doc.core_properties.keywords = 'Qbook, facility booking, approval, Qhazanah Sabah Berhad'
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == '__main__':
    build_document()
