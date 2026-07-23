from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from pathlib import Path
import os

OUT = Path('deliverables') / 'Qbook_Internal_Facility_Booking_Management_System_Proposal_Final.docx'
ASSETS = Path('deliverables') / 'proposal-assets'

NAVY = '0B2545'
BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
MUTED = '5B6573'
LIGHT_BLUE = 'E8EEF5'
LIGHT_GRAY = 'F4F6F9'
BORDER = 'C9D3DE'
WHITE = 'FFFFFF'
BLACK = '000000'

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = f'w:{edge}'
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ['val', 'sz', 'space', 'color']:
                if key in kwargs[edge]:
                    element.set(qn(f'w:{key}'), str(kwargs[edge][key]))

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')

def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)

def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement('w:keepNext')
    p_pr.append(keep)

def set_run_font(run, size=None, color=BLACK, bold=None, italic=None):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)

def setup_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in [
        ('Heading 1', 16, BLUE, 18, 10),
        ('Heading 2', 13, BLUE, 12, 6),
        ('Heading 3', 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    toc = doc.styles.add_style('TOC Entry', WD_STYLE_TYPE.PARAGRAPH)
    toc.font.name = 'Calibri'
    toc._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    toc._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    toc.font.size = Pt(11)
    toc.font.color.rgb = RGBColor.from_string(NAVY)
    toc.paragraph_format.space_after = Pt(5)
    toc.paragraph_format.line_spacing = 1.15

    label = doc.styles.add_style('Metadata Label', WD_STYLE_TYPE.PARAGRAPH)
    label.font.name = 'Calibri'
    label.font.size = Pt(10)
    label.font.bold = True
    label.font.color.rgb = RGBColor.from_string(MUTED)
    label.paragraph_format.space_after = Pt(1)

    small = doc.styles.add_style('Small Note', WD_STYLE_TYPE.PARAGRAPH)
    small.font.name = 'Calibri'
    small.font.size = Pt(9)
    small.font.color.rgb = RGBColor.from_string(MUTED)
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.0

def setup_section(section, first=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    if not first:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.paragraph_format.space_after = Pt(0)
        run = header.add_run('QBOOK | INTERNAL FACILITY BOOKING AND APPROVAL SYSTEM PROPOSAL')
        set_run_font(run, size=8.5, color=MUTED, bold=True)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.paragraph_format.space_before = Pt(0)
        footer.paragraph_format.space_after = Pt(0)
        run = footer.add_run('Qhazanah Sabah Berhad | Internal Use | Page ')
        set_run_font(run, size=8.5, color=MUTED)
        add_page_number(footer)

def add_paragraph(doc, text='', bold_prefix=None, after=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run_font(r)
    return p

def add_section_heading(doc, number, title):
    p = doc.add_paragraph(style='Heading 1')
    prefix = f'{number}. ' if number else ''
    r = p.add_run(f'{prefix}{title}')
    set_run_font(r, size=16, color=BLUE, bold=True)
    set_keep_with_next(p)
    return p

def add_toc_entry(doc, number, title, page):
    p = doc.add_paragraph(style='TOC Entry')
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    prefix = f'{number}. ' if number else ''
    r = p.add_run(f'{prefix}{title}')
    set_run_font(r, size=11, color=NAVY)
    r = p.add_run('\t' + str(page))
    set_run_font(r, size=11, color=NAVY)
    return p

def add_role_table(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [1900, 7460])
    table.style = 'Table Grid'
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, ['Role', 'Core responsibilities']):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=NAVY, bold=True)
    data = [
        ('Employee', [
            'Manage own bookings, including changes, rescheduling, and cancellations.',
            'Invite internal attendees and tag involved departments.',
            'View booking invitations, RSVP status, booking status, and calendar information.',
        ]),
        ('Admin', [
            'Review, approve, or reject bookings in accordance with defined requirements.',
            'Manage operational booking records and provide day-to-day support.',
            'Review booking activity, reports, and notification status.',
        ]),
        ('Super Admin', [
            'Manage users, roles, facilities, departments, department mailboxes, settings, reports, notifications, and audit records.',
        ]),
    ]
    for role, items in data:
        row = table.add_row()
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        row.cells[0].text = ''
        p = row.cells[0].paragraphs[0]
        r = p.add_run(role)
        set_run_font(r, size=10.5, color=NAVY, bold=True)
        row.cells[1].text = ''
        for idx, item in enumerate(items):
            p = row.cells[1].paragraphs[0] if idx == 0 else row.cells[1].add_paragraph()
            p.style = 'List Bullet'
            p.paragraph_format.left_indent = Inches(0.24)
            p.paragraph_format.first_line_indent = Inches(-0.16)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(item)
            set_run_font(r, size=9.7)
        for cell in row.cells:
            set_cell_border(cell, top={'val':'single','sz':'4','color':BORDER}, bottom={'val':'single','sz':'4','color':BORDER}, left={'val':'single','sz':'4','color':BORDER}, right={'val':'single','sz':'4','color':BORDER})
    return table

def add_pilot_table(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2500, 6860])
    table.style = 'Table Grid'
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, ['Pilot area', 'Proposed scope']):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ''
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, size=10.5, color=NAVY, bold=True)
    rows = [
        ('Participants', 'IT/Admin representatives and one selected department, with a defined group of internal staff users.'),
        ('Booking workflow', 'Create bookings; check availability; add department tags; invite internal attendees; submit for approval; approve or reject; issue cancellations.'),
        ('Communication checks', 'Confirm booking confirmations, reminders, approval, rejection, invitation and cancellation emails reach staff and relevant department mailboxes.'),
        ('Acceptance outcome', 'Record user feedback, confirm operating responsibilities, resolve pilot findings and obtain acceptance sign-off before recommending wider rollout.'),
    ]
    for left, right in rows:
        row = table.add_row()
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for cell, txt, bold in [(row.cells[0], left, True), (row.cells[1], right, False)]:
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            set_run_font(r, size=10, color=NAVY if bold else BLACK, bold=bold)
            set_cell_border(cell, top={'val':'single','sz':'4','color':BORDER}, bottom={'val':'single','sz':'4','color':BORDER}, left={'val':'single','sz':'4','color':BORDER}, right={'val':'single','sz':'4','color':BORDER})
    return table

def add_approval_table(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2250, 2150, 2200, 2760])
    table.style = 'Table Grid'
    headers = ['Approval item', 'Name', 'Signature', 'Date']
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ''
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, size=10, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    for role in ['Business sponsor / Management representative', 'IT owner / System owner', 'Pilot acceptance representative']:
        row = table.add_row()
        for i, cell in enumerate(row.cells):
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(14)
            r = p.add_run(role if i == 0 else '')
            set_run_font(r, size=9.6, color=NAVY if i == 0 else BLACK, bold=(i == 0))
            set_cell_border(cell, top={'val':'single','sz':'4','color':BORDER}, bottom={'val':'single','sz':'4','color':BORDER}, left={'val':'single','sz':'4','color':BORDER}, right={'val':'single','sz':'4','color':BORDER})
    return table

def add_image_with_caption(doc, image_name, caption_text):
    image_path = ASSETS / image_name
    if not image_path.exists():
        print(f"Warning: Image {image_name} not found at {image_path}")
        return
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(image_path), width=Inches(5.5))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    r_cap = p_cap.add_run(caption_text)
    set_run_font(r_cap, size=9.5, color=MUTED, italic=True)


def build_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    setup_section(doc.sections[0], first=True)

    # Cover Page
    for _ in range(4):
        doc.add_paragraph()
        
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('QHAZANAH SABAH BERHAD')
    set_run_font(r, size=13, color=MUTED, bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Qbook: Internal Facility\nBooking and Approval System')
    set_run_font(r, size=25, color=NAVY, bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('A proposed digital solution to improve internal facility booking management')
    set_run_font(r, size=13, color=MUTED, italic=True)
    
    p.paragraph_format.space_after = Pt(26)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(16)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    
    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [4680, 4680])
    for row in meta.rows:
        for cell in row.cells:
            set_cell_shading(cell, LIGHT_GRAY)
            set_cell_border(cell, top={'val':'single','sz':'4','color':WHITE}, bottom={'val':'single','sz':'4','color':WHITE}, left={'val':'single','sz':'4','color':WHITE}, right={'val':'single','sz':'4','color':WHITE})
            
    meta_data = [
        ('PREPARED FOR', 'Qhazanah Sabah Berhad'),
        ('PREPARED BY', 'Angelo Sugali Eddie, IT Intern'),
        ('DATE', '23 July 2026'),
        ('VERSION', '1.0'),
    ]
    for row_idx, (label, value) in enumerate(meta_data):
        cell_label = meta.cell(row_idx, 0)
        cell_value = meta.cell(row_idx, 1)
        
        cell_label.text = ''
        p = cell_label.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(label)
        set_run_font(r, size=8.5, color=MUTED, bold=True)
        
        cell_value.text = ''
        p = cell_value.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=NAVY, bold=True)
        
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    r = p.add_run('Submitted for consideration and stakeholder feedback')
    set_run_font(r, size=10.5, color=MUTED, italic=True)

    doc.add_page_break()
    
    # Table of Contents
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('TABLE OF CONTENTS')
    set_run_font(r, size=18, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(18)
    
    toc_entries = [
        ('1', 'Executive Summary', '3'),
        ('2', 'Background and Current Challenges', '3'),
        ('3', 'Proposed System: Qbook', '4'),
        ('4', 'Interface Overview', '6'),
        ('5', 'How the System Works', '8'),
        ('6', 'User Roles and System Access', '8'),
        ('7', 'Administration and Flexibility', '9'),
        ('8', 'Security and Governance', '9'),
        ('9', 'Potential Benefits', '10'),
        ('10', 'Proposed Pilot Rollout', '10'),
        ('11', 'Information Required if a Pilot Is Considered', '11'),
        ('12', 'Conclusion and Requested Next Step', '11'),
        ('', 'Appendix A. Technical Overview', '12'),
        ('', 'Appendix B. Approval and Sign-Off', '12'),
    ]
    for number, title, page in toc_entries:
        add_toc_entry(doc, number, title, page)

    doc.add_page_break()
    
    # Section 1
    add_section_heading(doc, '1', 'Executive Summary')
    add_paragraph(doc, 'Qbook is an internally developed web-based system for managing bookings of meeting rooms, event spaces and other designated internal facilities. It provides a single process for a staff member to submit a booking request, verify availability, identify attendees and relevant departments, obtain approval where required, and receive the related booking communications.')
    add_paragraph(doc, 'For authorised administrators, Qbook provides an operational record of booking status, approvals, notifications and activity. The system is proposed as a structured improvement to the current Excel-based process, not as a statement that it has already been approved, adopted or deployed across the organisation.')
    add_paragraph(doc, 'The system has been developed and tested during the current internship period. A working version is available for demonstration and review through a secure test deployment. If there is interest, Qbook can be prepared for a controlled pilot involving IT/Admin and one selected department before any wider rollout decision is made.')
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), LIGHT_GRAY)
    p_pr.append(shd)
    r = p.add_run('Proposed value — One controlled record of the booking request, availability, approval decision, notifications and subsequent changes.')
    set_run_font(r, size=11, color=NAVY, bold=True)

    # Section 2
    add_section_heading(doc, '2', 'Background and Current Challenges')
    add_paragraph(doc, 'Excel is currently used to manage facility bookings. Spreadsheets remain useful for recording information, but the process relies on manual checking and coordination when booking activity involves multiple facilities, changing schedules, attendees, departments and approval requirements. The following limitations can arise in that environment:')
    
    for item in [
        'Manual checking of facility availability.',
        'Difficulty maintaining one reliable source of booking information.',
        'Risk of overlapping or conflicting bookings.',
        'Manual coordination of approvals, attendees and departments.',
        'Limited visibility of booking changes and history.',
        'Difficulty tracking notifications and operational follow-up.',
    ]:
        add_bullet(doc, item)
        
    add_paragraph(doc, 'These are workflow limitations rather than shortcomings of individual staff or departments. Qbook is designed to provide a clearer operational process around the same underlying booking responsibilities.')

    # Section 3
    add_section_heading(doc, '3', 'Proposed System: Qbook')
    add_paragraph(doc, 'Qbook is a central internal facility booking system for meeting rooms, event spaces and other designated facilities. A booking record can capture the booking purpose, facility, date and time, attendee count, internal participants, relevant departments and supporting operational details. The system then applies the applicable availability and approval rules before the booking is treated as confirmed.')
    
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2340, 7020])
    table.style = 'Table Grid'
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, ['Capability', 'How it supports operations']):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ''
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, size=10.5, color=NAVY, bold=True)
        
    capabilities = [
        ('Booking management', 'Record the facility, date, time, purpose, attendee count and required details; allow authorised users to amend, reschedule or cancel eligible bookings.'),
        ('Availability control', 'Check the selected facility against active bookings and prevent overlapping bookings for the same time period. Back-to-back bookings are permitted.'),
        ('Collaboration and notices', 'Invite internal attendees, record RSVP responses and associate relevant departments so the people and internal units involved receive the appropriate booking updates.'),
        ('Approvals', 'Route selected facility bookings for review and record the approval, rejection or confirmation outcome before operational follow-up.'),
        ('Department coordination', 'Tag one or more departments for each booking. Notify department mailboxes about relevant booking updates. Display department involvement on booking details, approval forms and notifications.'),
        ('Notifications and email', 'Send templated email notifications for booking confirmations, reminders, approvals, rejections, invitations and cancellations. Users can manage their notification preferences.'),
        ('Microsoft 365 coordination', 'When configured, create Outlook calendar events for confirmed bookings, update them when bookings change and remove them when bookings are cancelled.'),
        ('Reporting and exports', 'Provide booking history, facility utilisation, cancellation, user activity and audit reports. Export records as CSV for further analysis.'),
        ('Operational oversight', 'Provide a booking calendar, printable approval form, email queue monitoring, notification status and activity history for authorised operational users.'),
    ]
    for left, right in capabilities:
        row = table.add_row()
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for cell, txt, bold in [(row.cells[0], left, True), (row.cells[1], right, False)]:
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            set_run_font(r, size=10, color=NAVY if bold else BLACK, bold=bold)
            set_cell_border(cell, top={'val':'single','sz':'4','color':BORDER}, bottom={'val':'single','sz':'4','color':BORDER}, left={'val':'single','sz':'4','color':BORDER}, right={'val':'single','sz':'4','color':BORDER})
            
    p = doc.add_paragraph(style='Heading 2')
    r = p.add_run('Microsoft 365 Calendar Integration')
    set_run_font(r, size=13, color=BLUE, bold=True)
    
    add_paragraph(doc, 'Qbook supports Microsoft company-account sign-in for controlled internal access. For teams that already use Outlook to coordinate meetings, Qbook is designed to complement that familiar calendar practice: Qbook remains the controlled source of the booking request, availability check, approval decision and operational history, while Microsoft 365 presents the confirmed event in Outlook.')
    add_paragraph(doc, 'When the calendar integration is configured, a confirmed Qbook booking creates an Outlook calendar event. A reschedule or other booking update updates the corresponding Outlook event; a cancellation removes it. The organisation may later select the intended destination: a central booking calendar, a facility or room calendar, or the booking owner\'s company calendar.')
    add_paragraph(doc, 'Qbook manages internal attendee invitations and RSVP responses. If the organisation chooses to send Outlook meeting invitations as well, attendee synchronisation can be configured as part of a later adoption decision. This preserves a clear invitation process and avoids duplicate notices by default.')

    doc.add_page_break()
    
    # Section 4
    add_section_heading(doc, '4', 'Interface Overview')
    add_paragraph(doc, 'The following screens illustrate how employees and administrators interact with Qbook. The system features a responsive employee-facing interface for booking activities and a comprehensive admin console for operational management.')
    
    p = doc.add_paragraph(style='Heading 2')
    r = p.add_run('Employee Interface')
    set_run_font(r, size=13, color=BLUE, bold=True)
    
    add_image_with_caption(doc, 'login.png', 'Figure 1. Secure Microsoft company-account sign-in.')
    add_image_with_caption(doc, 'qbook-dashboard.png', 'Figure 2. Employee dashboard and quick actions.')
    add_image_with_caption(doc, 'qbook-create-booking.png', 'Figure 3. Create booking: facility selection and availability.')
    add_image_with_caption(doc, 'calendar.png', 'Figure 4. Booking calendar: month view with status and visibility controls.')
    
    doc.add_page_break()
    
    p = doc.add_paragraph(style='Heading 2')
    r = p.add_run('Admin Console')
    set_run_font(r, size=13, color=BLUE, bold=True)
    
    add_image_with_caption(doc, 'admin dash.png', 'Figure 5. Admin console: operational dashboard with sidebar navigation.')
    add_image_with_caption(doc, 'approval page.png', 'Figure 6. Pending approvals queue: review, approve or reject booking requests.')
    add_image_with_caption(doc, 'facilities.png', 'Figure 7. Facility management: configure rooms, capacity, approval requirements and equipment.')
    add_image_with_caption(doc, 'reports.png', 'Figure 8. Reports and exports: booking metrics, utilisation data and CSV export controls.')

    doc.add_page_break()
    
    # Section 5
    add_section_heading(doc, '5', 'How the System Works')
    add_paragraph(doc, 'The proposed workflow gives employees a defined booking path while preserving the review, approval and operational visibility required by authorised administrators.')
    
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3000, 6360])
    table.style = 'Table Grid'
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, ['Step', 'Description']):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ''
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, size=10.5, color=NAVY, bold=True)
        
    steps = [
        ('1 | Create booking', 'Select facility, date and time.'),
        ('2 | Coordinate', 'Add attendees and departments where needed.'),
        ('3 | Review', 'Submit for approval when required.'),
        ('4 | Confirm', 'Receive confirmation and notifications.'),
        ('5 | Manage', 'Update invitations, reschedule or cancel.'),
    ]
    for left, right in steps:
        row = table.add_row()
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for cell, txt, bold in [(row.cells[0], left, True), (row.cells[1], right, False)]:
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            set_run_font(r, size=10, color=NAVY if bold else BLACK, bold=bold)
            set_cell_border(cell, top={'val':'single','sz':'4','color':BORDER}, bottom={'val':'single','sz':'4','color':BORDER}, left={'val':'single','sz':'4','color':BORDER}, right={'val':'single','sz':'4','color':BORDER})

    # Section 6
    add_section_heading(doc, '6', 'User Roles and System Access')
    add_paragraph(doc, 'The following role model separates ordinary employee activity from operational review and higher-risk system administration. Final role assignments would remain subject to organisational review if adoption is considered.')
    add_role_table(doc)
    
    # Section 7
    add_section_heading(doc, '7', 'Administration and Flexibility')
    add_paragraph(doc, 'The system is designed to be maintained through an Admin console. Subject to future organisational review, authorised administrators could maintain the following operating records and settings:')
    for item in [
        'Facilities and their relevant booking requirements.',
        'Users, roles and access status.',
        'Departments and associated mailboxes.',
        'Approval requirements for selected facilities.',
        'Equipment assignments for facilities.',
        'Blocked periods and maintenance schedules.',
        'Booking records, notification status, reports and activity history.',
        'System settings and integration configuration.',
    ]:
        add_bullet(doc, item)
    add_paragraph(doc, 'This structure allows the system to start with an agreed set of facilities and operating rules, then be adjusted as facilities, departments and operational requirements change. New facilities, departments and users can be added through the Admin console without requiring development changes.')

    # Section 8
    add_section_heading(doc, '8', 'Security and Governance')
    add_paragraph(doc, 'Qbook is designed with security controls appropriate for an internal company service. The following measures are built into the system:')
    for item in [
        'Microsoft company-email sign-in for user authentication. No separate passwords are stored by Qbook.',
        'Pre-provisioned, allowlisted access only. There is no public self-registration; user accounts must be created by an authorised administrator.',
        'Role-based access control enforced at the application and database layers. Employees, Admins and Super Admins see and manage only the functions appropriate to their assigned role.',
        'Database-level booking conflict prevention. The system prevents overlapping bookings for the same facility and time period through database constraints, not just user-interface checks.',
        'Approval workflow controls for facilities that require review before confirmation.',
        'An audit trail for important booking actions, including booking creation, status changes, approval decisions and administrative activity.',
        'Controlled access to administrative functions. Super Admin responsibility covers higher-risk configuration and governance tasks.',
        'Automated testing and continuous integration. Code changes are verified through automated tests before deployment.',
    ]:
        add_bullet(doc, item)
    add_paragraph(doc, 'These controls are designed to support accountability and operational continuity. Detailed integration settings, including Microsoft login, email and calendar configuration, would be confirmed with the IT team as part of pilot preparation.')

    # Section 9
    add_section_heading(doc, '9', 'Potential Benefits')
    for item in [
        'A single view of facility availability and booking status for authorised users.',
        'Reduced reliance on manual cross-checking of spreadsheet entries and separate messages.',
        'Fewer avoidable booking conflicts through controlled, database-enforced availability checks.',
        'A documented approval path for facilities that require review before confirmation.',
        'Consistent, automated communications to booking owners, invited staff and relevant departments.',
        'Traceable records of booking changes, notifications and authorised administrative activity.',
        'CSV-exportable reports for booking history, facility utilisation, cancellation rates and user activity.',
        'A practical foundation for additional facilities and departments if future adoption proceeds.',
    ]:
        add_bullet(doc, item)

    # Section 10
    add_section_heading(doc, '10', 'Proposed Pilot Rollout')
    add_paragraph(doc, 'It is recommended that Qbook first be introduced through a controlled pilot involving IT/Admin and one selected department. The pilot should use real internal booking scenarios within an agreed, limited facility scope. This will allow the company to validate the process, communications and role responsibilities before considering wider use.')
    add_pilot_table(doc)

    # Section 11
    add_section_heading(doc, '11', 'Information Required if a Pilot Is Considered')
    add_paragraph(doc, 'The following would be required only if stakeholders agree to explore a controlled pilot:')
    for item in [
        'Named system owner, Super Admins and Admins.',
        'Existing department mailboxes.',
        'Initial staff list: full name, company email, department and assigned role.',
        'Confirmation of facilities to include at launch.',
        'Approved Qbook sender mailbox/address for booking-related emails.',
        'Approval to send booking-related emails to staff and department mailboxes.',
        'Pilot participants and acceptance sign-off representatives.',
    ]:
        add_bullet(doc, item)

    # Section 12
    add_section_heading(doc, '12', 'Conclusion and Requested Next Step')
    add_paragraph(doc, 'Qbook is presented as a practical internal alternative to an Excel-based facility booking process. It provides a defined record of each booking, its availability outcome, approval status, communications and subsequent changes, while retaining Outlook as a familiar coordination surface where Microsoft 365 calendar integration is selected.')
    add_paragraph(doc, 'The system has been developed with attention to security, data integrity and operational controls. A working version is available for demonstration and review.')
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), LIGHT_GRAY)
    p_pr.append(shd)
    r = p.add_run('Requested next step — This proposal is submitted for consideration and stakeholder feedback. Stakeholders are invited to indicate whether Qbook should progress to a controlled pilot discussion. Only if there is interest would the relevant owners subsequently assess operational requirements, system access and implementation arrangements.')
    set_run_font(r, size=11, color=NAVY, bold=True)

    doc.add_page_break()

    # Appendix A
    p = doc.add_paragraph(style='Heading 1')
    r = p.add_run('Appendix A. Technical Overview')
    set_run_font(r, size=16, color=BLUE, bold=True)
    set_keep_with_next(p)
    
    p = doc.add_paragraph(style='Heading 2')
    r = p.add_run('System Architecture')
    set_run_font(r, size=13, color=BLUE, bold=True)
    
    add_paragraph(doc, 'Qbook is built with a modern web technology stack designed for reliability, security and maintainability. The following summarises the technical foundation:')
    
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2340, 7020])
    table.style = 'Table Grid'
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, ['Area', 'Detail']):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ''
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, size=10.5, color=NAVY, bold=True)
        
    tech_details = [
        ('Frontend', 'Next.js (React) with TypeScript. Responsive design using Tailwind CSS.'),
        ('Backend', 'Next.js server-side rendering and API routes. Server Actions for secure data operations.'),
        ('Database', 'PostgreSQL (managed by Supabase). Row-Level Security (RLS) policies enforce access control at the data layer.'),
        ('Authentication', 'Microsoft Entra (Azure AD) OAuth via Supabase Auth. Pre-provisioned access with company domain restriction.'),
        ('Email', 'Configurable email provider (Resend or SMTP). Templated notifications with queue management and retry logic.'),
        ('Calendar sync', 'Optional one-way outbound synchronisation to Microsoft 365 (Graph API) or n8n webhook.'),
        ('Data integrity', 'Database exclusion constraints prevent overlapping bookings. Atomic transactions for booking creation with participants and departments.'),
        ('Testing', 'Automated unit and integration tests (Vitest). End-to-end browser tests (Playwright). Secret scanning and security checks.'),
        ('CI/CD', 'GitHub Actions for continuous integration. Automated linting, type checking, testing and production build verification.'),
        ('Deployment', 'A Vercel-hosted demonstration environment is available. Any company-designated deployment would require IT review, environment configuration, integration permissions and operational acceptance.'),
        ('Source control', 'Version-controlled with Git (GitHub). Full commit history with 160+ commits documenting development progress.'),
    ]
    for left, right in tech_details:
        row = table.add_row()
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for cell, txt, bold in [(row.cells[0], left, True), (row.cells[1], right, False)]:
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            set_run_font(r, size=10, color=NAVY if bold else BLACK, bold=bold)
            set_cell_border(cell, top={'val':'single','sz':'4','color':BORDER}, bottom={'val':'single','sz':'4','color':BORDER}, left={'val':'single','sz':'4','color':BORDER}, right={'val':'single','sz':'4','color':BORDER})

    p = doc.add_paragraph(style='Heading 2')
    r = p.add_run('Current System Metrics')
    set_run_font(r, size=13, color=BLUE, bold=True)
    
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3000, 6360])
    table.style = 'Table Grid'
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, ['Metric', 'Value']):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ''
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, size=10.5, color=NAVY, bold=True)
        
    metrics = [
        ('Database tables', '25 tables defined through version-controlled migrations (Row-Level Security enabled on all tables; deployed schema subject to environment verification)'),
        ('Database migrations', '39 (append-only, version-controlled)'),
        ('Automated test suites', '47 unit/integration tests + 7 end-to-end browser test suites'),
        ('Application components', '90+ React components across employee and admin interfaces'),
        ('API endpoints', '6 secure server endpoints for booking operations, notifications and integrations'),
        ('Documentation', '28 internal technical documents covering architecture, deployment, security and operations'),
        ('Development period', 'April – July 2026 (160+ commits)'),
    ]
    for left, right in metrics:
        row = table.add_row()
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for cell, txt, bold in [(row.cells[0], left, True), (row.cells[1], right, False)]:
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            set_run_font(r, size=10, color=NAVY if bold else BLACK, bold=bold)
            set_cell_border(cell, top={'val':'single','sz':'4','color':BORDER}, bottom={'val':'single','sz':'4','color':BORDER}, left={'val':'single','sz':'4','color':BORDER}, right={'val':'single','sz':'4','color':BORDER})

    # Appendix B
    p = doc.add_paragraph(style='Heading 1')
    r = p.add_run('Appendix B. Approval and Sign-Off')
    set_run_font(r, size=16, color=BLUE, bold=True)
    set_keep_with_next(p)
    
    add_paragraph(doc, 'The signatories below acknowledge review of this proposal and the decision on whether to proceed with the controlled pilot.')
    add_approval_table(doc)

    doc.core_properties.title = 'Qbook: Internal Facility Booking and Approval System Proposal'
    doc.core_properties.subject = 'Internal proposal for Qbook facility booking system'
    doc.core_properties.author = 'Angelo Sugali Eddie'
    doc.core_properties.keywords = 'Qbook, facility booking, approval, Qhazanah Sabah Berhad'
    doc.save(OUT)
    print(f"Generated {OUT.resolve()}")

if __name__ == '__main__':
    build_document()
